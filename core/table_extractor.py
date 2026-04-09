"""Extract and analyze table structures from Word documents."""
import logging
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
from docx import Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)


class TableExtractor:
    """Extract tables from Word documents with structure preservation."""

    def __init__(self):
        """Initialize table extractor."""
        self.logger = logger

    def extract_all_tables(self, docx_path: str) -> List[Dict[str, Any]]:
        """
        Extract all tables from Word document.

        Args:
            docx_path: Path to .docx file

        Returns:
            List of table dictionaries with structure and metadata
        """
        try:
            doc = Document(docx_path)
            tables = []

            for table_idx, table in enumerate(doc.tables):
                table_data = self._extract_table(table, table_idx)
                tables.append(table_data)

            self.logger.info(f"Extracted {len(tables)} tables from {docx_path}")
            return tables

        except Exception as e:
            self.logger.error(f"Error extracting tables: {e}")
            return []

    def _extract_table(self, table: Table, table_idx: int) -> Dict[str, Any]:
        """
        Extract single table structure.

        Args:
            table: docx Table object
            table_idx: Table index in document

        Returns:
            Dictionary with table structure
        """
        rows = []

        for row_idx, row in enumerate(table.rows):
            row_data = self._extract_row(row, row_idx)
            rows.append(row_data)

        return {
            "table_idx": table_idx,
            "row_count": len(table.rows),
            "col_count": len(table.columns),
            "rows": rows,
            "metadata": {
                "style": table.style.name if table.style else "Normal",
                "autofit": table.autofit if hasattr(table, 'autofit') else None,
            }
        }

    def _extract_row(self, row, row_idx: int) -> Dict[str, Any]:
        """
        Extract row data with merged cell info.

        Args:
            row: docx Row object
            row_idx: Row index

        Returns:
            Row data with cell information
        """
        cells = []

        for col_idx, cell in enumerate(row.cells):
            cell_data = self._extract_cell(cell, row_idx, col_idx)
            cells.append(cell_data)

        return {
            "row_idx": row_idx,
            "cells": cells,
            "height": row.height.pt if row.height else None,
        }

    def _extract_cell(self, cell: _Cell, row_idx: int, col_idx: int) -> Dict[str, Any]:
        """
        Extract cell content and metadata.

        Args:
            cell: docx Cell object
            row_idx: Row index
            col_idx: Column index

        Returns:
            Cell data with content and formatting
        """
        text_content = cell.text.strip()

        # Check for merged cells
        tcPr = cell._element.tcPr
        is_merged = False
        merge_info = None

        if tcPr is not None:
            # Check for vertical merge
            vMerge = tcPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vMerge')
            hMerge = tcPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridSpan')

            if vMerge is not None or hMerge is not None:
                is_merged = True
                merge_info = {
                    "vertical_merge": vMerge is not None,
                    "horizontal_span": int(hMerge.get('w:val', 1)) if hMerge is not None else 1,
                }

        return {
            "row_idx": row_idx,
            "col_idx": col_idx,
            "text": text_content,
            "is_merged": is_merged,
            "merge_info": merge_info,
            "paragraphs": len(cell.paragraphs),
            "shading": self._get_cell_shading(cell),
        }

    def _get_cell_shading(self, cell: _Cell) -> Optional[str]:
        """Get cell background color if any."""
        try:
            tcPr = cell._element.tcPr
            if tcPr is None:
                return None

            shd = tcPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
            if shd is not None:
                return shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
            return None
        except Exception as e:
            self.logger.debug(f"Error getting cell shading: {e}")
            return None

    def find_table_by_content(self, tables: List[Dict], search_text: str) -> Optional[Tuple[int, int, int]]:
        """
        Find table and cell containing specific text.

        Args:
            tables: List of extracted tables
            search_text: Text to search for

        Returns:
            Tuple of (table_idx, row_idx, col_idx) or None
        """
        search_lower = search_text.lower()

        for table_idx, table in enumerate(tables):
            for row_idx, row in enumerate(table.get("rows", [])):
                for col_idx, cell in enumerate(row.get("cells", [])):
                    if search_lower in cell.get("text", "").lower():
                        return (table_idx, row_idx, col_idx)

        return None

    def get_table_context(self, table: Dict, row_idx: int, col_idx: int,
                         context_rows: int = 2) -> Dict[str, Any]:
        """
        Get context around a cell (header + surrounding rows).

        Args:
            table: Table dictionary
            row_idx: Row index
            col_idx: Column index
            context_rows: Number of rows above/below to include

        Returns:
            Context dictionary with header and surrounding rows
        """
        rows = table.get("rows", [])

        # Get header row (first row)
        header_row = rows[0] if rows else None

        # Get context rows
        start_row = max(0, row_idx - context_rows)
        end_row = min(len(rows), row_idx + context_rows + 1)
        context = rows[start_row:end_row]

        return {
            "header": header_row,
            "target_row_idx": row_idx,
            "target_col_idx": col_idx,
            "context_rows": context,
            "row_range": (start_row, end_row),
        }
