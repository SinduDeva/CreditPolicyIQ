"""Word document handler module for CreditPolicyIQ."""
import logging
from typing import Dict, List, Any, Optional
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

logger = logging.getLogger(__name__)


class DocxHandler:
    """Handler for Word document operations."""

    def __init__(self):
        """Initialize DOCX handler."""
        self.logger = logger
        self._structure_cache: Optional[Dict[str, Any]] = None
        self._current_doc: Optional[Document] = None
        self._current_path: Optional[str] = None

    def extract_structure(self, docx_path: str) -> Dict[str, Any]:
        """
        Extract document structure from DOCX file.

        Args:
            docx_path: Path to DOCX file

        Returns:
            Dictionary with sections, paragraphs, and all_text
        """
        if not Path(docx_path).exists():
            self.logger.error(f"DOCX file not found: {docx_path}")
            return {"sections": [], "paragraphs": [], "all_text": "", "error": "File not found"}

        try:
            doc = Document(docx_path)
            self._current_doc = doc
            self._current_path = docx_path

            paragraphs_list = []
            sections = []
            current_section = None
            all_text = ""

            # Iterate through paragraphs
            for para_idx, para in enumerate(doc.paragraphs):
                para_data = {
                    "index": para_idx,
                    "text": para.text,
                    "style": para.style.name if para.style else "Normal",
                    "level": para.paragraph_format.outline_level
                    if para.paragraph_format.outline_level is not None
                    else 0,
                    "is_heading": self._is_heading(para),
                }

                paragraphs_list.append(para_data)
                all_text += para.text + "\n"

                # Track sections by heading
                if para_data["is_heading"]:
                    if current_section:
                        sections.append(current_section)
                    current_section = {
                        "title": para.text,
                        "start_para_idx": para_idx,
                        "paragraphs": [],
                    }
                elif current_section:
                    current_section["paragraphs"].append(para_idx)

            # Add last section
            if current_section:
                sections.append(current_section)

            # Cache structure
            structure = {
                "sections": sections,
                "paragraphs": paragraphs_list,
                "all_text": all_text,
                "total_paragraphs": len(doc.paragraphs),
            }
            self._structure_cache = structure

            self.logger.info(
                f"Extracted structure from {docx_path}: {len(sections)} sections, "
                f"{len(paragraphs_list)} paragraphs"
            )
            return structure

        except Exception as e:
            self.logger.error(f"Error extracting structure: {e}")
            return {"sections": [], "paragraphs": [], "all_text": "", "error": str(e)}

    def get_paragraph_context(
        self, para_index: int, context_lines: int = 2
    ) -> Dict[str, Any]:
        """
        Get paragraph context (before, current, after).

        Args:
            para_index: Index of paragraph
            context_lines: Number of lines before/after to include

        Returns:
            Dictionary with before_text, current_text, after_text
        """
        if not self._current_doc:
            self.logger.warning("No document loaded")
            return {"before": [], "current": None, "after": [], "error": "No document"}

        try:
            paragraphs = self._current_doc.paragraphs

            if para_index < 0 or para_index >= len(paragraphs):
                self.logger.warning(f"Invalid paragraph index: {para_index}")
                return {"before": [], "current": None, "after": [], "error": "Invalid index"}

            # Get context
            before_start = max(0, para_index - context_lines)
            after_end = min(len(paragraphs), para_index + context_lines + 1)

            before_texts = [paragraphs[i].text for i in range(before_start, para_index)]
            current_text = paragraphs[para_index].text
            after_texts = [
                paragraphs[i].text for i in range(para_index + 1, after_end)
            ]

            return {
                "before": before_texts,
                "current": current_text,
                "after": after_texts,
                "para_index": para_index,
            }

        except Exception as e:
            self.logger.error(f"Error getting paragraph context: {e}")
            return {"before": [], "current": None, "after": [], "error": str(e)}

    def update_paragraph(
        self, para_index: int, new_text: str, output_path: str
    ) -> bool:
        """
        Update paragraph text while preserving formatting.

        Args:
            para_index: Index of paragraph to update
            new_text: New paragraph text
            output_path: Path to save updated document

        Returns:
            True if successful, False otherwise
        """
        if not self._current_doc:
            self.logger.error("No document loaded")
            return False

        try:
            paragraphs = self._current_doc.paragraphs

            if para_index < 0 or para_index >= len(paragraphs):
                self.logger.error(f"Invalid paragraph index: {para_index}")
                return False

            para = paragraphs[para_index]

            # Clear existing runs and add new text with same style
            for run in para.runs:
                run._element.getparent().remove(run._element)

            new_run = para.add_run(new_text)
            if para.runs and len(para.runs) > 1:
                # Copy style from first remaining run
                first_run = para.runs[0]
                new_run.font.name = first_run.font.name
                new_run.font.size = first_run.font.size
                new_run.bold = first_run.bold
                new_run.italic = first_run.italic

            self._current_doc.save(output_path)
            self.logger.info(f"Updated paragraph {para_index} and saved to {output_path}")
            return True

        except Exception as e:
            self.logger.error(f"Error updating paragraph: {e}")
            return False

    def add_paragraph_after(
        self, para_index: int, new_text: str, style: str = "Normal"
    ) -> Optional[int]:
        """
        Add new paragraph after specified paragraph.

        Args:
            para_index: Index of paragraph after which to insert
            new_text: Text for new paragraph
            style: Paragraph style name

        Returns:
            Index of new paragraph or None if failed
        """
        if not self._current_doc:
            self.logger.error("No document loaded")
            return None

        try:
            paragraphs = self._current_doc.paragraphs

            if para_index < 0 or para_index >= len(paragraphs):
                self.logger.error(f"Invalid paragraph index: {para_index}")
                return None

            # Insert paragraph after the specified one
            para = paragraphs[para_index]
            new_para = para.insert_paragraph_before(new_text)
            if style:
                new_para.style = style

            self.logger.info(
                f"Added new paragraph after index {para_index} with text: {new_text[:50]}..."
            )
            return para_index + 1

        except Exception as e:
            self.logger.error(f"Error adding paragraph: {e}")
            return None

    def _is_heading(self, paragraph) -> bool:
        """
        Check if paragraph is a heading.

        Args:
            paragraph: docx paragraph object

        Returns:
            True if heading, False otherwise
        """
        if not paragraph.style:
            return False
        style_name = paragraph.style.name.lower()
        return "heading" in style_name
